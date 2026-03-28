"""
report_service.py
Generates a professionally formatted Word (.docx) report with:
  - Cover page
  - Descriptive statistics table
  - Pearson correlation matrix table
  - Regression results tables (one per model)
  - Diagnostic tests summary table
  - AI-written narrative discussion
"""

import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────────────────────────
# Style constants
# ─────────────────────────────────────────────

HEADER_COLOR = RGBColor(0x18, 0x5F, 0xA5)   # Deep blue
ROW_ALT_COLOR = RGBColor(0xF0, 0xF4, 0xF8)  # Light blue-grey
SIG_COLOR = RGBColor(0x0F, 0x6E, 0x56)      # Teal for significance
BODY_FONT = "Cambria"
HEADING_FONT = "Calibri"


# ─────────────────────────────────────────────
# Document helpers
# ─────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    """Set cell background colour via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _cell_text(cell, text: str, bold=False, font_size=10,
               color: RGBColor = None, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = BODY_FONT
    if color:
        run.font.color.rgb = color


def _add_header_row(table, headers: list[str], col_widths: list[float] = None):
    """Style the first row as a header."""
    row = table.rows[0]
    for i, (cell, text) in enumerate(zip(row.cells, headers)):
        _set_cell_bg(cell, "185FA5")
        _cell_text(cell, text, bold=True, font_size=10,
                   color=RGBColor(0xFF, 0xFF, 0xFF))
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for i, width in enumerate(col_widths):
            table.columns[i].width = Inches(width)


def _zebra_row(row, idx: int):
    if idx % 2 == 0:
        for cell in row.cells:
            _set_cell_bg(cell, "F0F4F8")


def _add_table_note(doc: Document, note: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"Note: {note}")
    run.italic = True
    run.font.size = Pt(9)
    run.font.name = BODY_FONT
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)


def _sig_stars(p_value) -> str:
    if p_value is None:
        return ""
    if p_value < 0.01:
        return "***"
    if p_value < 0.05:
        return "**"
    if p_value < 0.10:
        return "*"
    return ""


def _fmt(val, decimals=4) -> str:
    if val is None:
        return "—"
    try:
        return f"{float(val):.{decimals}f}"
    except Exception:
        return str(val)


def _fmt_pval(val) -> str:
    if val is None:
        return "—"
    f = float(val)
    if f < 0.001:
        return "<0.001"
    return f"{f:.4f}"


# ─────────────────────────────────────────────
# Sections
# ─────────────────────────────────────────────

def _add_cover_page(doc: Document, title: str, dependent_var: str,
                    independent_vars: list[str]):
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PANEL DATA REGRESSION ANALYSIS")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = HEADING_FONT
    run.font.color.rgb = HEADER_COLOR

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = HEADING_FONT

    doc.add_paragraph()

    for label, value in [
        ("Dependent Variable", dependent_var),
        ("Independent Variables", ", ".join(independent_vars)),
        ("Report Generated", datetime.now().strftime("%B %d, %Y")),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}: ")
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = BODY_FONT
        run2 = p.add_run(value)
        run2.font.size = Pt(12)
        run2.font.name = BODY_FONT

    doc.add_page_break()


def _add_descriptive_stats_table(doc: Document, descriptive_stats: dict):
    doc.add_heading("1. Descriptive Statistics", level=1)

    headers = ["Variable", "N", "Mean", "Std. Dev.", "Min", "25th Pct.",
               "Median", "75th Pct.", "Max", "Skewness", "Kurtosis"]
    col_widths = [1.2, 0.5, 0.85, 0.85, 0.75, 0.75, 0.75, 0.75, 0.75, 0.85, 0.85]

    table = doc.add_table(rows=1 + len(descriptive_stats), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    _add_header_row(table, headers, col_widths)

    for i, (var, s) in enumerate(descriptive_stats.items(), start=1):
        row = table.rows[i]
        _zebra_row(row, i)
        values = [
            var,
            _fmt(s.get("count"), 0),
            _fmt(s.get("mean")),
            _fmt(s.get("std")),
            _fmt(s.get("min")),
            _fmt(s.get("p25")),
            _fmt(s.get("median")),
            _fmt(s.get("p75")),
            _fmt(s.get("max")),
            _fmt(s.get("skewness")),
            _fmt(s.get("kurtosis")),
        ]
        for j, (cell, val) in enumerate(zip(row.cells, values)):
            align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            _cell_text(cell, val, font_size=10, align=align)

    doc.add_paragraph()
    _add_table_note(doc, "Descriptive statistics computed for all study variables.")


def _add_correlation_table(doc: Document, correlation_matrix: dict):
    doc.add_heading("2. Pearson Correlation Matrix", level=1)

    variables = correlation_matrix.get("variables", [])
    coeffs = correlation_matrix.get("coefficients", {})
    pvals = correlation_matrix.get("pvalues", {})
    n = len(variables)

    headers = ["Variable"] + [f"({i+1})" for i in range(n)]
    table = doc.add_table(rows=1 + n + 1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    _add_header_row(table, headers)

    # Variable labels row
    label_row = table.rows[1 + n]
    label_row.cells[0].text = ""
    for i, var in enumerate(variables):
        _cell_text(label_row.cells[i + 1], f"({i+1}) {var}",
                   font_size=9, align=WD_ALIGN_PARAGRAPH.LEFT)

    for i, var in enumerate(variables):
        row = table.rows[i + 1]
        _zebra_row(row, i)
        _cell_text(row.cells[0], f"({i+1}) {var}",
                   bold=True, font_size=10, align=WD_ALIGN_PARAGRAPH.LEFT)

        for j, var2 in enumerate(variables):
            cell = row.cells[j + 1]
            if i == j:
                _cell_text(cell, "1.000", font_size=10)
            else:
                coeff = coeffs.get(var, {}).get(var2)
                pval = pvals.get(var, {}).get(var2)
                stars = _sig_stars(pval)
                val_str = f"{_fmt(coeff, 3)}{stars}"
                color = SIG_COLOR if stars else None
                _cell_text(cell, val_str, font_size=10, color=color)

    doc.add_paragraph()
    _add_table_note(
        doc,
        "*** p<0.01, ** p<0.05, * p<0.10. Pearson correlation coefficients reported."
    )


def _add_regression_table(doc: Document, regression_results: dict):
    doc.add_heading("3. Panel Regression Results", level=1)

    model_keys = list(regression_results.keys())
    model_objects = [regression_results[k] for k in model_keys]
    model_labels = [m.get("model", k) for k, m in zip(model_keys, model_objects)]

    # Collect all variables (union)
    all_vars: list[str] = []
    for m in model_objects:
        for v in m.get("variables", []):
            if v["variable"] not in all_vars:
                all_vars.append(v["variable"])

    # One sub-table per model with full stats
    for model_key, model_data in regression_results.items():
        model_name = model_data.get("model", model_key)
        doc.add_heading(f"3.{list(regression_results.keys()).index(model_key)+1}  {model_name}", level=2)

        headers = ["Variable", "β (Coef.)", "Std. Error", "t-Statistic",
                   "p-Value", "Sig.", "LCL (95%)", "UCL (95%)"]
        col_widths = [1.3, 0.9, 0.9, 0.9, 0.9, 0.5, 0.9, 0.9]

        vars_data = model_data.get("variables", [])
        table = doc.add_table(rows=1 + len(vars_data), cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _add_header_row(table, headers, col_widths)

        for i, vd in enumerate(vars_data, start=1):
            row = table.rows[i]
            _zebra_row(row, i)
            pv = vd.get("p_value")
            stars = _sig_stars(pv)
            text_color = SIG_COLOR if stars else None

            row_vals = [
                (vd.get("variable"), True, WD_ALIGN_PARAGRAPH.LEFT, None),
                (_fmt(vd.get("beta")), False, WD_ALIGN_PARAGRAPH.CENTER, text_color),
                (_fmt(vd.get("std_error")), False, WD_ALIGN_PARAGRAPH.CENTER, None),
                (_fmt(vd.get("t_stat")), False, WD_ALIGN_PARAGRAPH.CENTER, None),
                (_fmt_pval(pv), False, WD_ALIGN_PARAGRAPH.CENTER, text_color),
                (stars, True, WD_ALIGN_PARAGRAPH.CENTER, text_color),
                (_fmt(vd.get("lcl")), False, WD_ALIGN_PARAGRAPH.CENTER, None),
                (_fmt(vd.get("ucl")), False, WD_ALIGN_PARAGRAPH.CENTER, None),
            ]
            for j, (cell, (val, bold, align, color)) in enumerate(
                zip(row.cells, row_vals)
            ):
                _cell_text(cell, val, bold=bold, font_size=10,
                           align=align, color=color)

        # Model fit statistics below table
        doc.add_paragraph()
        p = doc.add_paragraph()
        stats_parts = []
        for label, key in [
            ("R²", "r_squared"),
            ("R² (within)", "r_squared_within"),
            ("R² (between)", "r_squared_between"),
            ("R² (overall)", "r_squared_overall"),
            ("F-stat", "f_statistic"),
            ("F p-value", "f_pvalue"),
            ("Observations", "n_obs"),
            ("Entities", "n_entities"),
            ("Periods", "n_periods"),
        ]:
            val = model_data.get(key)
            if val is not None:
                stats_parts.append(f"{label} = {_fmt(val, 4) if isinstance(val, float) else val}")
        run = p.add_run("  |  ".join(stats_parts))
        run.font.size = Pt(9)
        run.font.name = BODY_FONT
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

        _add_table_note(
            doc,
            "*** p<0.01, ** p<0.05, * p<0.10. LCL/UCL are 95% confidence interval bounds. "
            "Robust standard errors used where applicable."
        )
        doc.add_paragraph()


def _add_diagnostics_table(doc: Document, diagnostic_results: dict):
    doc.add_heading("4. Diagnostic Tests", level=1)

    # Summary table
    summary_rows = []

    h = diagnostic_results.get("hausman")
    if h and "error" not in h:
        summary_rows.append([
            "Hausman Test",
            f"χ² = {_fmt(h.get('statistic'))}",
            f"df = {h.get('df')}",
            _fmt_pval(h.get("p_value")),
            h.get("preferred_model", "—"),
        ])

    bp = diagnostic_results.get("breusch_pagan")
    if bp and "error" not in bp:
        detected = "Detected" if bp.get("heteroskedasticity_detected") else "Not detected"
        summary_rows.append([
            "Breusch-Pagan Test",
            f"LM = {_fmt(bp.get('lm_statistic'))}",
            "—",
            _fmt_pval(bp.get("lm_p_value")),
            f"Heteroskedasticity: {detected}",
        ])

    vif = diagnostic_results.get("vif")
    if vif and "error" not in vif:
        concern = "Yes (VIF > 10)" if vif.get("multicollinearity_concern") else "No"
        summary_rows.append([
            "VIF (Multicollinearity)",
            f"Max VIF = {_fmt(vif.get('max_vif'))}",
            "—",
            "—",
            f"Concern: {concern}",
        ])

    wld = diagnostic_results.get("wooldridge")
    if wld and "error" not in wld:
        detected = "Detected" if wld.get("serial_correlation_detected") else "Not detected"
        summary_rows.append([
            "Wooldridge Test",
            f"t = {_fmt(wld.get('t_statistic'))}",
            f"df = {wld.get('df', '—')}",
            _fmt_pval(wld.get("p_value")),
            f"Serial correlation: {detected}",
        ])

    if summary_rows:
        headers = ["Test", "Statistic", "df", "p-Value", "Conclusion"]
        table = doc.add_table(rows=1 + len(summary_rows), cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _add_header_row(table, headers, [1.6, 1.3, 0.6, 0.8, 1.9])

        for i, row_data in enumerate(summary_rows, start=1):
            row = table.rows[i]
            _zebra_row(row, i)
            for j, (cell, val) in enumerate(zip(row.cells, row_data)):
                align = WD_ALIGN_PARAGRAPH.LEFT if j in (0, 4) else WD_ALIGN_PARAGRAPH.CENTER
                _cell_text(cell, val, font_size=10, align=align)

    # VIF detail table
    if vif and "values" in vif:
        doc.add_paragraph()
        doc.add_heading("4.1  VIF Detail", level=2)
        vif_vals = [v for v in vif["values"] if v["variable"] != "const"]
        vtable = doc.add_table(rows=1 + len(vif_vals), cols=3)
        vtable.style = "Table Grid"
        vtable.alignment = WD_TABLE_ALIGNMENT.CENTER
        _add_header_row(vtable, ["Variable", "VIF", "Assessment"])

        for i, vv in enumerate(vif_vals, start=1):
            row = vtable.rows[i]
            _zebra_row(row, i)
            v = vv["vif"]
            assessment = "Severe" if v > 10 else "Moderate" if v > 5 else "Acceptable"
            color = (RGBColor(0xA3, 0x2D, 0x2D) if v > 10
                     else RGBColor(0xBA, 0x75, 0x17) if v > 5 else None)
            _cell_text(row.cells[0], vv["variable"], bold=True,
                       align=WD_ALIGN_PARAGRAPH.LEFT)
            _cell_text(row.cells[1], _fmt(v))
            _cell_text(row.cells[2], assessment, color=color)

    doc.add_paragraph()
    _add_table_note(
        doc,
        "VIF < 5 = no concern; 5–10 = moderate; > 10 = severe. "
        "Breusch-Pagan tests H₀ of homoskedasticity. "
        "Wooldridge tests H₀ of no first-order serial correlation."
    )


def _add_narrative(doc: Document, narrative: str):
    doc.add_page_break()
    doc.add_heading("5. Results and Discussion", level=1)

    if not narrative:
        doc.add_paragraph("AI narrative not generated.")
        return

    for line in narrative.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        else:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.size = Pt(11)
                run.font.name = BODY_FONT


# ─────────────────────────────────────────────
# Master builder
# ─────────────────────────────────────────────

def build_word_report(
    title: str,
    dependent_var: str,
    independent_vars: list[str],
    descriptive_stats: dict,
    correlation_matrix: dict,
    regression_results: dict,
    diagnostic_results: dict,
    llm_narrative: str = "",
) -> bytes:
    """
    Assembles the complete Word report and returns it as bytes.
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default paragraph style
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(11)

    _add_cover_page(doc, title, dependent_var, independent_vars)
    _add_descriptive_stats_table(doc, descriptive_stats)
    doc.add_paragraph()
    _add_correlation_table(doc, correlation_matrix)
    doc.add_paragraph()
    _add_regression_table(doc, regression_results)
    _add_diagnostics_table(doc, diagnostic_results)
    _add_narrative(doc, llm_narrative)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
